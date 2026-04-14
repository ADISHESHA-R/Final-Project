"""Extract plain text from .docx (Office Open XML)."""

from __future__ import annotations

import re
from io import BytesIO

from docx import Document


def extract_docx_text(data: bytes, max_chars: int = 120_000) -> str:
    doc = Document(BytesIO(data))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = (cell.text or "").strip()
                if t:
                    parts.append(t)
    raw = "\n".join(parts)
    raw = re.sub(r"\s+", " ", raw).strip()
    if len(raw) > max_chars:
        raw = raw[:max_chars]
    return raw
